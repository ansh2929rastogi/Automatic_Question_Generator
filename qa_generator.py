from model_loader import tokenizer, model, device
import torch
import re
import random


# ---------- UTILITIES ----------

def word_count(text):
    return len(text.split())


def auto_question_count(summary):
    """
    Adaptive question count based on summary length
    """

    wc = word_count(summary)

    if wc < 150:
        return 4
    elif wc < 300:
        return 6
    elif wc < 600:
        return 10
    else:
        return 14


def split_sentences(text):
    text = text.replace("\n", " ").strip()
    sentences = re.split(r'(?<=[.!?])\s+', text)

    # Keep only meaningful sentences
    sentences = [s.strip() for s in sentences if len(s.strip()) > 60]

    return sentences if sentences else [text]


# ---------- QUESTION GENERATION ----------

def generate_question(sentence):
    """
    Generates descriptive academic questions only
    """

    prompt = (
        "Generate one clear, well-formed academic comprehension question based on the following text. "
        "The question should test understanding of concepts, definitions, differences, causes, or purposes. "
        "Do not use multiple-choice style and do not use 'which of the following'.\n\n"
        f"Text:\n{sentence}\n\nQuestion:"
    )

    inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=384).to(device)

    outputs = model.generate(
        **inputs,
        max_length=64,
        num_beams=3,
        early_stopping=True,
        no_repeat_ngram_size=3,
        repetition_penalty=2.0
    )

    question = tokenizer.decode(outputs[0], skip_special_tokens=True).strip()

    if not question.endswith("?"):
        question += "?"

    return question


# ---------- QUALITY FILTER ----------

def is_bad_question(q):

    if len(q) < 30:
        return True

    banned = [
        "which of the following",
        "choose the correct",
        "select the correct",
        "pick the correct",
        "identify the option"
    ]

    for b in banned:
        if b in q.lower():
            return True

    return False


# ---------- MAIN PIPELINE (QUESTIONS ONLY) ----------

def generate_qa_pairs(summary: str, _ignored=None):

    # Decide number of questions automatically
    target_questions = auto_question_count(summary)

    sentences = split_sentences(summary)
    random.shuffle(sentences)

    results = []
    used_questions = set()

    attempts = 0
    max_attempts = target_questions * 6   # Allows filling quota without slowing too much

    while len(results) < target_questions and attempts < max_attempts:
        attempts += 1
        sent = random.choice(sentences)

        try:
            question = generate_question(sent)

            if is_bad_question(question):
                continue

            if question.lower() in used_questions:
                continue

            used_questions.add(question.lower())

            results.append({
                "question": question
            })

        except:
            continue

    return results


# ---------- DOCX EXPORT (QUESTIONS ONLY) ----------

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def save_qa_to_docx(qa_pairs, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    title = doc.add_heading("Generated Questions", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, qa in enumerate(qa_pairs, 1):
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{idx}. {qa['question']}")
        q_run.bold = True
        q_run.font.size = Pt(13)

        doc.add_paragraph("")

    doc.save(output_path)
